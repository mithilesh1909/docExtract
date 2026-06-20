import pdfplumber
import docx
from PIL import Image
import pytesseract
import io
from pathlib import Path


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract raw text from uploaded file based on extension."""
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(file_bytes)
    elif ext == ".docx":
        return _extract_docx(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")
    elif ext in {".png", ".jpg", ".jpeg"}:
        return _extract_image(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def _extract_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs).strip()


def _extract_image(file_bytes: bytes) -> str:
    try:
        image = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(image).strip()
    except Exception:
        return ""
